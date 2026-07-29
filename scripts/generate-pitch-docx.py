#!/usr/bin/env python3
"""Generate BaytMiftah investor pitch document (DOCX)."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "BaytMiftah_Pitch.docx"

BRAND_GREEN = RGBColor(0x0F, 0x29, 0x22)
BRAND_ORANGE = RGBColor(0xD4, 0x6A, 0x4A)
MUTED = RGBColor(0x6B, 0x72, 0x80)


def set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("BaytMiftah")
    run.bold = True
    run.font.size = Pt(36)
    run.font.color.rgb = BRAND_GREEN

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("بيت مفتاح  ·  The Real Estate Operating System for West Africa")
    r.font.size = Pt(14)
    r.font.color.rgb = MUTED

    doc.add_paragraph()
    tag = doc.add_paragraph()
    tag.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = tag.add_run("Investor & Partner Pitch  ·  July 2026")
    tr.font.size = Pt(12)
    tr.italic = True

    link = doc.add_paragraph()
    link.alignment = WD_ALIGN_PARAGRAPH.CENTER
    lr = link.add_run("www.baytmiftah.com")
    lr.font.color.rgb = BRAND_ORANGE
    lr.bold = True

    doc.add_page_break()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = BRAND_GREEN


def add_bullet(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        b = p.add_run(bold_prefix)
        b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
        for p in hdr[i].paragraphs:
            for run in p.runs:
                run.bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
    doc.add_paragraph()


def build_pitch() -> Document:
    doc = Document()
    set_doc_defaults(doc)
    add_title_page(doc)

    # Executive summary
    add_heading(doc, "Executive Summary")
    doc.add_paragraph(
        "BaytMiftah is a Ghana-first Real Estate Operating System (REOS) that unifies property "
        "discovery, transactions, and professional operations in one platform. Unlike fragmented "
        "listing sites, BaytMiftah serves consumers (buy, rent, lease, short stay), landlords and "
        "agencies (CRM, payments, team workspace), and platform operators (trust, moderation, compliance) "
        "from a single product — already live at baytmiftah.com."
    )
    doc.add_paragraph(
        "Mission: Give every person and property business in emerging markets a trusted, end-to-end "
        "digital home — from discovery to keys in hand."
    )
    doc.add_paragraph(
        "We are building the default property infrastructure for West Africa, expanding to Nigeria, "
        "Francophone West Africa, East Africa, and global diaspora markets."
    )

    # Problem
    add_heading(doc, "The Problem")
    for item in [
        ("Fragmented tools ", "— agents juggle WhatsApp, spreadsheets, and separate payment apps"),
        ("No end-to-end journeys ", "— renters, buyers, and hosts use different platforms for each step"),
        ("Trust gap ", "— fraud, unverified listings, and manual KYC slow transactions"),
        ("Payments friction ", "— rent deposits, escrow, and cross-border cards are poorly integrated"),
        ("Agencies lack ERP ", "— small and mid-size firms cannot afford enterprise software"),
    ]:
        add_bullet(doc, item[1], item[0])

    # Solution
    add_heading(doc, "Our Solution")
    doc.add_paragraph(
        "BaytMiftah is one operating system for the full property lifecycle:"
    )
    add_table(
        doc,
        ["Audience", "What they get"],
        [
            ["Consumers", "Search, save, message, book stays, apply, pay, manage leases & maintenance"],
            ["Landlords & agencies", "Multi-org workspace: listings, CRM, calendar, finance, team, AI"],
            ["Platform", "Admin console: moderation, trust/KYC, fraud alerts, support"],
        ],
    )

    # Product highlights
    add_heading(doc, "Product Highlights (Live Today)")
    highlights = [
        "Public marketplace with search, maps, compare, and verified agencies",
        "Consumer hub: saved properties, messaging, applications, viewings, wallet",
        "Short-stay booking + host workspace (trips, reservations, calendar)",
        "Tenant portal: leases, rent payments, maintenance requests",
        "Agency workspace: listings, leads, contacts, tasks, documents, automation",
        "Paystack payments live; escrow & wallet MVP; Paystack webhooks deployed",
        "12+ languages; location-based market onboarding",
        "Mobile PWA + Capacitor native wrappers (Android/iOS)",
        "Role-optimized UI for consumers, hosts, agents, analysts, and admins",
        "BaytMiftah AI integrated on home, search, property, consumer, and workspace (works without API keys)",
    ]
    for h in highlights:
        add_bullet(doc, h)

    # Market
    add_heading(doc, "Market Opportunity")
    doc.add_paragraph(
        "West Africa's real estate market is large, informal, and rapidly digitizing. Ghana alone "
        "has growing urban rental demand, a rising short-stay segment, and an underserved agency "
        "sector that needs affordable operations software."
    )
    add_table(
        doc,
        ["Segment", "Opportunity"],
        [
            ["Residential rental & lease", "Recurring rent flows, tenant lifecycle, maintenance"],
            ["Short-stay / hospitality", "Airbnb-style inventory managed by local hosts"],
            ["Agency operations", "CRM, lead pipeline, team collaboration, compliance"],
            ["Adjacent finance", "Mortgage, insurance, escrow — high-margin add-ons"],
        ],
    )

    # International expansion
    add_heading(doc, "International Expansion")
    add_table(
        doc,
        ["Wave", "Markets", "Timeline"],
        [
            ["Wave 1", "Ghana (live)", "Now"],
            ["Wave 2", "Nigeria, Côte d'Ivoire, Senegal", "6–12 months"],
            ["Wave 3", "Kenya, South Africa", "12–18 months"],
            ["Wave 4", "UK, US, UAE diaspora", "18–24 months"],
        ],
    )
    doc.add_paragraph(
        "Each market launch includes localized payments (Paystack/Stripe), currency, language, "
        "agency onboarding, and legal compliance. Multi-currency and 12+ language support are already built."
    )

    # AI
    add_heading(doc, "BaytMiftah AI")
    doc.add_paragraph(
        "AI is integrated across the product today. Guided help (FAQ, natural-language search parsing, "
        "listing templates) works without any API keys. Adding OPENAI_API_KEY upgrades to smart mode "
        "for conversational assistance, better search, and listing copy generation."
    )
    add_bullet(doc, "Home, search, property detail, consumer dashboard, workspace")
    add_bullet(doc, "Edge functions: parse-search-query + ai-assistant")
    add_bullet(doc, "Future: lead scoring, pricing AI, AWS Bedrock for regional models")

    # AWS migration
    add_heading(doc, "Infrastructure Roadmap (AWS)")
    doc.add_paragraph(
        "Current stack (Vercel + Supabase) is optimal for launch. As we scale past ~50k MAU or sign "
        "enterprise contracts requiring VPC/compliance, we will progressively migrate to AWS:"
    )
    for item in [
        "Phase 1: S3 media, Lambda webhooks, SES email (hybrid with Supabase)",
        "Phase 2: Aurora PostgreSQL, OpenSearch, regional read replicas",
        "Phase 3: Full AWS multi-region with Bedrock AI and Cognito auth",
    ]:
        add_bullet(doc, item)

    # Business model
    add_heading(doc, "Business Model")
    add_table(
        doc,
        ["Revenue stream", "Description"],
        [
            ["Pro workspace subscriptions", "Per-seat or per-org SaaS for agencies and landlords"],
            ["Enterprise white-label", "Custom branding, multi-org insights, workflows"],
            ["Transaction fees", "Payment processing margin on rent, deposits, bookings"],
            ["Financial services", "Mortgage & insurance marketplace commissions (planned)"],
            ["API & integrations", "Developer platform, MLS syndication, partner webhooks"],
        ],
    )

    # Traction
    add_heading(doc, "Traction & Status")
    add_table(
        doc,
        ["Milestone", "Status"],
        [
            ["Production launch", "Live at www.baytmiftah.com (Vercel + Supabase)"],
            ["Payments", "Paystack integrated; Stripe wired, keys pending"],
            ["Test coverage", "58 automated tests passing"],
            ["Module readiness", "Marketplace 90% · Consumer 90% · Agency ERP 80%"],
            ["Mobile", "PWA live; native builds scaffolded"],
        ],
    )

    # Technology
    add_heading(doc, "Technology")
    doc.add_paragraph(
        "Modern, scalable stack built for speed and low operating cost:"
    )
    add_bullet(doc, "React 18, TypeScript, Vite — fast web and mobile UI")
    add_bullet(doc, "Supabase — PostgreSQL, auth, realtime, storage, edge functions")
    add_bullet(doc, "Vercel — global CDN hosting")
    add_bullet(doc, "Paystack + Stripe — local and international payments")
    add_bullet(doc, "Capacitor — single codebase for Android and iOS")

    # Roadmap
    add_heading(doc, "Roadmap")
    add_heading(doc, "Near term (2–4 weeks)", level=2)
    for item in [
        "Resend email (confirmations, receipts, password reset)",
        "Stripe live for international cards",
        "Google & Apple OAuth",
        "KYC verification completion",
        "Subscription billing (Pro / Enterprise)",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Medium term (1–3 months)", level=2)
    for item in [
        "Mortgage & insurance marketplaces",
        "Vendor marketplace for maintenance",
        "Trust & reputation system",
        "AI pricing, lead scoring, maintenance triage",
        "App Store & Play Store launch",
        "MLS listing syndication",
    ]:
        add_bullet(doc, item)

    add_heading(doc, "Long-term vision", level=2)
    doc.add_paragraph(
        "Become the default REOS for West Africa, then expand internationally — covering discover, "
        "transact, operate, trust, and monetize across the full property value chain."
    )

    # Competitive advantage
    add_heading(doc, "Why BaytMiftah Wins")
    for item in [
        ("Full-stack REOS ", "— not just listings; operations, payments, and trust in one product"),
        ("Ghana-first payments ", "— Paystack-native, mobile-money ready, escrow built in"),
        ("Role-aware product ", "— consumer and pro experiences optimized per user type"),
        ("Multi-tenant workspace ", "— agencies, teams, and permissions out of the box"),
        ("International-ready ", "— 12+ locales, multi-currency, Stripe for global cards"),
        ("Ship velocity ", "— production platform built and deployed; iterating weekly"),
    ]:
        add_bullet(doc, item[1], item[0])

    # The ask
    add_heading(doc, "The Ask")
    doc.add_paragraph(
        "We are seeking strategic partners and investment to:"
    )
    for item in [
        "Accelerate go-to-market in Ghana and West Africa",
        "Complete payment, email, and OAuth integrations",
        "Launch native mobile apps on App Store and Play Store",
        "Scale agency onboarding and enterprise white-label sales",
        "Build mortgage, insurance, and vendor marketplace partnerships",
    ]:
        add_bullet(doc, item)

    doc.add_paragraph()
    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cr = contact.add_run("Contact: support@baytmiftah.com  ·  www.baytmiftah.com")
    cr.bold = True
    cr.font.color.rgb = BRAND_GREEN

    # Appendix
    doc.add_page_break()
    add_heading(doc, "Appendix: Module Completion")
    add_table(
        doc,
        ["Module", "Completion"],
        [
            ["Marketplace", "90%"],
            ["Consumer experience", "90%"],
            ["Agent CRM", "85%"],
            ["Agency ERP (workspace)", "80%"],
            ["Property management", "75%"],
            ["Smart property", "65%"],
            ["Financial services", "60%"],
            ["Developer platform", "70%"],
            ["Enterprise platform", "60%"],
            ["Trust & compliance", "80%"],
            ["Mobile platform", "90%"],
        ],
    )

    add_heading(doc, "Appendix: Key Integrations", level=2)
    add_table(
        doc,
        ["Service", "Purpose", "Status"],
        [
            ["Supabase", "Database, auth, storage", "Live"],
            ["Vercel", "Hosting", "Live"],
            ["Paystack", "Africa payments", "Live"],
            ["Stripe", "International cards", "Wired"],
            ["Resend", "Email", "Pending keys"],
            ["Google Maps", "Property maps", "Live"],
            ["OpenAI", "AI assistant", "Integrated (keys optional)"],
            ["Blockchain", "On-chain verification", "Deferred"],
        ],
    )

    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = "BaytMiftah — Confidential  ·  July 2026"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc = build_pitch()
    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
